from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_NAME = "Sistema IRIS - EPSI HL"
TODAY = date.today().strftime("%d/%m/%Y")
OUTPUT_DIR = Path(__file__).resolve().parent / "manuales"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
        shade_cell(header_cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value))
    document.add_paragraph("")
    return table


def set_margins(section):
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_cover(document, title, subtitle):
    section = document.sections[0]
    set_margins(section)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(PROJECT_NAME)
    run.bold = True
    run.font.size = Pt(22)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(19)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(120)
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(11)

    add_table(
        document,
        ["Campo", "Valor"],
        [
            ["Organizacion", "EPSI HL"],
            ["Proyecto", PROJECT_NAME],
            ["Version del documento", "1.0"],
            ["Fecha de elaboracion", TODAY],
            ["Estado", "Aprobado para entrega documental"],
        ],
    )

    document.add_page_break()


def set_default_font(document):
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Title"].font.name = "Calibri"
    styles["Heading 1"].font.name = "Calibri"
    styles["Heading 2"].font.name = "Calibri"
    styles["Heading 3"].font.name = "Calibri"


def add_heading(document, text, level=1):
    document.add_heading(text, level=level)


def add_paragraph(document, text, bold_prefix=None):
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix, rest = text.split(":", 1)
        run = paragraph.add_run(prefix + ":")
        run.bold = True
        paragraph.add_run(rest)
    else:
        paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(2)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.add_run(item)
        paragraph.paragraph_format.space_after = Pt(2)


def add_footer(section, text):
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = text


def create_technical_manual():
    document = Document()
    set_default_font(document)
    add_cover(
        document,
        "Manual Tecnico",
        "Documento de referencia para instalacion, arquitectura, operacion y mantenimiento del sistema.",
    )

    add_heading(document, "1. Introduccion", level=1)
    add_paragraph(
        document,
        "El presente manual tecnico describe la solucion Sistema IRIS implementada para EPSI HL, su arquitectura general, componentes funcionales, configuracion, despliegue, seguridad, operacion y mantenimiento. El objetivo es proporcionar una base formal para administradores, desarrolladores y personal tecnico responsable del soporte del sistema."
    )

    add_heading(document, "2. Objetivo del sistema", level=1)
    add_paragraph(
        document,
        "Sistema IRIS centraliza la gestion de remisiones en PDF, clientes y usuarios con control de acceso por roles. La plataforma permite estandarizar procesos internos, mantener trazabilidad operativa y reducir la dispersion de informacion."
    )

    add_heading(document, "3. Alcance funcional", level=1)
    add_table(
        document,
        ["Modulo", "Descripcion", "Estado actual"],
        [
            ["Autenticacion", "Inicio de sesion, validacion de sesion y recuperacion de contrasena.", "Operativo"],
            ["Remisiones", "Creacion, consulta, edicion restringida y generacion de PDF.", "Operativo"],
            ["Clientes", "Consulta, registro, actualizacion y exportacion de datos.", "Operativo"],
            ["Usuarios", "Administracion de usuarios, roles, estados y reseteo de claves.", "Operativo"],
            ["Turnos", "Modulo proyectado para futura implementacion.", "Pendiente"],
            ["Reportes y BI", "Modulo proyectado para visualizacion ejecutiva.", "Pendiente"],
        ],
    )

    add_heading(document, "4. Arquitectura de la solucion", level=1)
    add_paragraph(
        document,
        "La solucion esta estructurada como una aplicacion web con separacion entre cliente, API, base de datos y capa de publicacion. El acceso de usuarios se realiza por medio de una interfaz web unificada, mientras que el procesamiento de negocio y persistencia de datos se delega al backend y a la base de datos."
    )
    add_bullets(
        document,
        [
            "Cliente web para autenticacion, navegacion y operacion de modulos.",
            "API central para reglas de negocio, validaciones, seguridad y generacion documental.",
            "Base de datos relacional para usuarios, clientes y remisiones.",
            "Proxy inverso para publicar una sola entrada a la aplicacion y enrutar solicitudes.",
            "Volumen persistente para almacenamiento de PDFs generados.",
        ],
    )

    add_heading(document, "5. Componentes principales", level=1)
    add_table(
        document,
        ["Componente", "Responsabilidad principal"],
        [
            ["Frontend", "Presentacion de la interfaz, flujos del usuario y consumo de servicios."],
            ["Backend", "Exposicion de servicios, autenticacion, autorizacion, validacion y generacion de remisiones."],
            ["Base de datos", "Persistencia de usuarios, clientes, remisiones y datos relacionados."],
            ["Proxy", "Publicacion de la aplicacion y enrutamiento hacia frontend, API, assets y healthcheck."],
            ["Servicio PDF", "Construccion del documento de remision segun plantilla institucional."],
        ],
    )

    add_heading(document, "6. Roles y permisos", level=1)
    add_table(
        document,
        ["Rol", "Capacidades principales"],
        [
            ["GERENCIAL", "Administracion de usuarios, consulta y edicion de remisiones, exportacion de clientes y anulacion de remisiones."],
            ["DIRECCION", "Administracion de usuarios y acceso general al sistema, sin exportacion de clientes ni edicion de remisiones."],
            ["SUPERVISION", "Operacion general autenticada sin privilegios administrativos."],
            ["ASISTENTE", "Operacion general autenticada sin privilegios administrativos."],
            ["APOYO", "Operacion general autenticada sin privilegios administrativos."],
            ["AUXILIARES", "Operacion general autenticada sin privilegios administrativos."],
        ],
    )

    add_heading(document, "7. Flujos tecnicos relevantes", level=1)
    add_numbered(
        document,
        [
            "Autenticacion del usuario mediante credenciales institucionales y emision de token de sesion.",
            "Consulta del perfil autenticado para aplicar permisos y habilitar u ocultar funcionalidades.",
            "Solicitud del consecutivo de remision y captura de informacion del cliente.",
            "Registro o actualizacion del cliente previo a la generacion de la remision.",
            "Construccion de la remision, calculo de totales y generacion del PDF institucional.",
            "Almacenamiento del documento generado y disponibilidad para consulta o descarga posterior.",
            "Gestion administrativa de usuarios con altas, bajas logicas, edicion y reseteo de contrasenas.",
        ],
    )

    add_heading(document, "8. Servicios expuestos por la API", level=1)
    add_table(
        document,
        ["Grupo", "Operaciones principales"],
        [
            ["Autenticacion", "Inicio de sesion, consulta de usuario actual y recuperacion de contrasena."],
            ["Usuarios", "Listado, creacion, actualizacion, eliminacion y reseteo de credenciales."],
            ["Clientes", "Consulta individual, registro/actualizacion y exportacion de informacion."],
            ["Remisiones", "Obtencion de consecutivo, creacion, consulta, edicion y descarga de PDF."],
            ["Salud del servicio", "Verificacion de disponibilidad operativa."],
        ],
    )

    add_heading(document, "9. Requisitos de implementacion", level=1)
    add_bullets(
        document,
        [
            "Entorno con acceso a la red corporativa o al servidor donde se despliegue la plataforma.",
            "Variables de configuracion definidas para seguridad, conexion de base de datos y publicacion.",
            "Base de datos disponible y accesible para la API.",
            "Servicio de correo configurado si se requiere recuperacion de contrasena por email.",
            "Puertos habilitados para acceso del proxy y conectividad entre servicios.",
        ],
    )

    add_heading(document, "10. Despliegue de la solucion", level=1)
    add_paragraph(
        document,
        "El proyecto contempla un despliegue unificado donde la base de datos, el backend, el frontend y el proxy pueden iniciarse de forma orquestada. La publicacion externa se realiza a traves del proxy, concentrando el acceso a la aplicacion en un unico punto de entrada."
    )
    add_bullets(
        document,
        [
            "La base de datos aloja la informacion operacional del sistema.",
            "La API ejecuta migraciones y publica los servicios de negocio.",
            "El frontend se compila y se entrega como sitio web estatico.",
            "El proxy distribuye trafico hacia la interfaz, la API, los assets y el healthcheck.",
        ],
    )

    add_heading(document, "11. Variables y configuracion critica", level=1)
    add_table(
        document,
        ["Parametro", "Finalidad"],
        [
            ["Secreto de autenticacion", "Firma y validacion de tokens de sesion."],
            ["Credenciales de base de datos", "Conexion de la API al motor relacional."],
            ["Contrasena inicial administrativa", "Provision del usuario administrador por defecto."],
            ["Origenes permitidos", "Control de acceso entre interfaz y API."],
            ["Ruta de salida de PDF", "Persistencia de documentos generados."],
            ["Configuracion SMTP", "Envio de solicitudes de recuperacion de contrasena."],
        ],
    )

    add_heading(document, "12. Seguridad", level=1)
    add_bullets(
        document,
        [
            "La aplicacion emplea autenticacion por token y validacion del estado activo del usuario en solicitudes protegidas.",
            "Se aplican restricciones de permisos por rol para operaciones administrativas y de consulta sensible.",
            "Existe limitacion de intentos de login para reducir abuso por direccion IP.",
            "Se incorporan cabeceras de seguridad a traves del backend.",
            "Las contrasenas deben mantenerse fuera de repositorios y archivos versionados.",
        ],
    )
    add_paragraph(
        document,
        "Como consideraciones de riesgo observables en la implementacion actual, se recomienda reemplazar credenciales por defecto antes de pasar a produccion, endurecer secretos de autenticacion, validar configuraciones de CORS y revisar el tratamiento de contrasenas temporales y visibles en procesos administrativos."
    )

    add_heading(document, "13. Base de datos y persistencia", level=1)
    add_paragraph(
        document,
        "El sistema utiliza una base de datos relacional como repositorio principal. Adicionalmente, mantiene compatibilidad de migracion desde una fuente historica previa. La estructura contempla tablas y relaciones asociadas a usuarios, clientes y remisiones, junto con scripts de migracion para inicializacion y traslado de datos."
    )

    add_heading(document, "14. Generacion documental", level=1)
    add_paragraph(
        document,
        "El modulo de remisiones genera documentos PDF con formato institucional. El proceso toma la informacion de la remision y del cliente, construye el archivo con la plantilla visual definida y lo deja disponible para descarga o consulta posterior."
    )

    add_heading(document, "15. Pruebas y control de calidad", level=1)
    add_bullets(
        document,
        [
            "El proyecto incorpora pruebas automatizadas para backend y frontend.",
            "Existen rutinas de cobertura para soporte a herramientas de calidad.",
            "Se recomienda ejecutar pruebas antes de desplegar cambios y validar los flujos criticos de autenticacion, usuarios, clientes y remisiones.",
        ],
    )

    add_heading(document, "16. Operacion y mantenimiento", level=1)
    add_bullets(
        document,
        [
            "Monitorear el estado del servicio mediante el endpoint de salud.",
            "Supervisar espacio de almacenamiento destinado a remisiones PDF.",
            "Mantener actualizadas las variables de entorno sensibles.",
            "Realizar respaldos periodicos de la base de datos.",
            "Registrar y revisar incidencias de acceso, generacion de PDF y conectividad.",
        ],
    )

    add_heading(document, "17. Incidencias comunes", level=1)
    add_table(
        document,
        ["Situacion", "Posible causa", "Accion recomendada"],
        [
            ["No es posible iniciar sesion", "Credenciales invalidas, usuario inactivo o configuracion incorrecta.", "Verificar correo institucional, estado del usuario y secreto de autenticacion."],
            ["No se genera la remision PDF", "Datos incompletos, error en persistencia o ruta de salida no disponible.", "Validar campos, almacenamiento y permisos del servicio."],
            ["No se envian correos de recuperacion", "SMTP no configurado o no disponible.", "Revisar parametros del correo y conectividad del servidor."],
            ["No carga la aplicacion desde el navegador", "Fallo en proxy, frontend o backend.", "Verificar publicacion del proxy, rutas y estado general del stack."],
        ],
    )

    add_heading(document, "18. Recomendaciones tecnicas", level=1)
    add_bullets(
        document,
        [
            "Mantener separados los valores de desarrollo, pruebas y produccion.",
            "Usar credenciales robustas y rotacion periodica de secretos.",
            "Documentar cambios estructurales en base de datos y servicios.",
            "Controlar el acceso administrativo a los modulos de usuarios y remisiones.",
            "Conservar este documento actualizado cuando cambie el alcance funcional del sistema.",
        ],
    )

    add_heading(document, "19. Conclusion", level=1)
    add_paragraph(
        document,
        "Sistema IRIS constituye una solucion integral para la operacion interna de EPSI HL en materia de remisiones, clientes y control de usuarios. Su documentacion tecnica permite facilitar la administracion del sistema, su despliegue, soporte y evolucion controlada."
    )

    for section in document.sections:
        add_footer(section, "Manual Tecnico - Sistema IRIS - EPSI HL")

    return document


def create_user_manual():
    document = Document()
    set_default_font(document)
    add_cover(
        document,
        "Manual de Usuario",
        "Guia operativa para el uso funcional de la plataforma por parte de usuarios finales.",
    )

    add_heading(document, "1. Presentacion", level=1)
    add_paragraph(
        document,
        "Este manual de usuario explica el funcionamiento general de Sistema IRIS y orienta a los usuarios en el uso de sus principales modulos. El documento esta dirigido al personal de EPSI HL que interactua con la plataforma para gestionar remisiones, clientes y usuarios de acuerdo con su perfil."
    )

    add_heading(document, "2. Objetivo", level=1)
    add_paragraph(
        document,
        "Brindar una guia clara y formal para que los usuarios puedan acceder al sistema, comprender su navegacion, ejecutar tareas frecuentes y resolver situaciones basicas durante la operacion diaria."
    )

    add_heading(document, "3. Alcance del sistema", level=1)
    add_bullets(
        document,
        [
            "Ingreso seguro al sistema mediante credenciales institucionales.",
            "Creacion de remisiones con generacion de documento PDF.",
            "Registro y consulta de informacion de clientes.",
            "Administracion de usuarios segun permisos autorizados.",
            "Consulta de modulos futuros visibles como referencia, aun no habilitados para operacion.",
        ],
    )

    add_heading(document, "4. Perfiles de usuario", level=1)
    add_table(
        document,
        ["Perfil", "Descripcion de uso"],
        [
            ["GERENCIAL", "Perfil administrativo con acceso a gestion de usuarios, exportacion de clientes y consulta/edicion de remisiones."],
            ["DIRECCION", "Perfil con acceso administrativo parcial para gestion de usuarios."],
            ["SUPERVISION", "Perfil operativo para uso general del sistema."],
            ["ASISTENTE", "Perfil operativo enfocado en procesos diarios."],
            ["APOYO", "Perfil operativo con acceso a funciones generales habilitadas."],
            ["AUXILIARES", "Perfil operativo con acceso a funciones generales habilitadas."],
        ],
    )

    add_heading(document, "5. Acceso al sistema", level=1)
    add_numbered(
        document,
        [
            "Abrir el enlace institucional del sistema en el navegador autorizado.",
            "Ingresar el correo institucional registrado y la contrasena asignada.",
            "Seleccionar la opcion de ingreso.",
            "Esperar la validacion de credenciales y la carga del panel principal.",
        ],
    )
    add_paragraph(
        document,
        "Si el usuario no recuerda su contrasena, puede utilizar la opcion de recuperacion disponible en la pantalla de acceso. El sistema solicitara el correo registrado y permitira aplicar la restauracion segun el procedimiento definido."
    )

    add_heading(document, "6. Pantalla principal", level=1)
    add_paragraph(
        document,
        "Una vez autenticado, el usuario accede al panel principal del sistema. Desde esta pantalla se visualizan los accesos a modulos, informacion general y las opciones habilitadas segun el rol asignado."
    )
    add_bullets(
        document,
        [
            "Inicio.",
            "Remisiones.",
            "Turnos del personal (referencia futura).",
            "Reportes (referencia futura).",
            "Inteligencia de negocio BI (referencia futura).",
            "Usuarios y contrasenas, solo para perfiles autorizados.",
        ],
    )

    add_heading(document, "7. Modulo de remisiones", level=1)
    add_paragraph(
        document,
        "El modulo de remisiones es el componente principal del sistema. Permite registrar la informacion del cliente, agregar items o conceptos, calcular valores y generar el documento PDF de forma estructurada."
    )

    add_heading(document, "7.1 Flujo general de remision", level=2)
    add_numbered(
        document,
        [
            "Ingresar al modulo Remisiones.",
            "Consultar el cliente por su identificacion o registrar sus datos si no existe.",
            "Completar la informacion general de la remision.",
            "Agregar los items correspondientes al servicio o producto.",
            "Revisar subtotal, impuestos y total.",
            "Confirmar el resumen final.",
            "Generar la remision PDF.",
        ],
    )

    add_heading(document, "7.2 Datos del cliente", level=2)
    add_bullets(
        document,
        [
            "Nombre o razon social.",
            "Documento o identificacion.",
            "Direccion.",
            "Ciudad.",
            "Telefono.",
            "Correo cuando aplique.",
        ],
    )
    add_paragraph(
        document,
        "El sistema puede recuperar informacion previa del cliente y permite actualizarla durante el proceso cuando corresponda."
    )

    add_heading(document, "7.3 Registro de items", level=2)
    add_bullets(
        document,
        [
            "Descripcion del item o concepto.",
            "Cantidad.",
            "Valor unitario.",
            "Subtotal.",
        ],
    )
    add_paragraph(
        document,
        "Cada item registrado se incorpora a la remision y contribuye al calculo del valor total del documento."
    )

    add_heading(document, "7.4 Consulta y edicion de remisiones", level=2)
    add_paragraph(
        document,
        "El perfil GERENCIAL puede consultar remisiones existentes, revisar su informacion, descargar nuevamente el PDF y realizar ajustes cuando la operacion lo requiera. Adicionalmente, este perfil puede marcar remisiones como anuladas."
    )

    add_heading(document, "8. Modulo de clientes", level=1)
    add_paragraph(
        document,
        "La informacion de clientes se gestiona desde el flujo de remisiones y mediante funcionalidades complementarias de consulta y actualizacion. El objetivo es mantener una base organizada para reutilizar datos en futuras operaciones."
    )
    add_bullets(
        document,
        [
            "Busqueda por numero de identificacion.",
            "Actualizacion de informacion basica.",
            "Persistencia de cambios para futuros documentos.",
            "Exportacion de clientes en perfiles autorizados.",
        ],
    )

    add_heading(document, "9. Modulo de usuarios y contrasenas", level=1)
    add_paragraph(
        document,
        "Los perfiles GERENCIAL y DIRECCION disponen del modulo de administracion de usuarios. Desde este espacio es posible crear cuentas, actualizar la informacion de acceso, activar o desactivar usuarios y ejecutar procesos de reseteo de contrasena."
    )

    add_heading(document, "9.1 Operaciones disponibles", level=2)
    add_bullets(
        document,
        [
            "Crear nuevo usuario.",
            "Editar datos de un usuario existente.",
            "Activar o desactivar usuarios.",
            "Eliminar usuarios segun autorizacion.",
            "Resetear contrasena temporal.",
        ],
    )

    add_heading(document, "10. Recuperacion y cambio de contrasena", level=1)
    add_numbered(
        document,
        [
            "Ubicar la opcion de recuperacion en la pantalla de acceso.",
            "Ingresar el correo institucional del usuario.",
            "Seguir las instrucciones entregadas por el sistema o el enlace recibido.",
            "Registrar una nueva contrasena valida.",
            "Ingresar nuevamente con la clave actualizada.",
        ],
    )

    add_heading(document, "11. Buenas practicas de uso", level=1)
    add_bullets(
        document,
        [
            "Verificar la informacion del cliente antes de generar la remision.",
            "Revisar cuidadosamente cantidades, valores y observaciones.",
            "No compartir credenciales con terceros.",
            "Cerrar sesion al finalizar la jornada o al dejar el equipo sin supervision.",
            "Solicitar soporte si se detectan errores de permisos o datos inconsistentes.",
        ],
    )

    add_heading(document, "12. Mensajes o situaciones frecuentes", level=1)
    add_table(
        document,
        ["Situacion", "Interpretacion", "Accion sugerida"],
        [
            ["Credenciales invalidas", "El correo o la contrasena no coinciden con un usuario activo.", "Intentar nuevamente o solicitar validacion al administrador."],
            ["Acceso denegado", "El perfil no tiene permiso para la operacion solicitada.", "Consultar con un usuario administrador o gerencial."],
            ["No se pudo generar la remision", "Existe informacion incompleta o una incidencia temporal del sistema.", "Revisar los campos obligatorios e intentar de nuevo."],
            ["No aparece un modulo", "El rol asignado no habilita esa funcionalidad.", "Confirmar permisos con el area responsable."],
        ],
    )

    add_heading(document, "13. Recomendaciones para administradores funcionales", level=1)
    add_bullets(
        document,
        [
            "Mantener actualizada la lista de usuarios activos e inactivos.",
            "Resetear contrasenas solo cuando exista validacion previa del solicitante.",
            "Restringir operaciones sensibles a personal autorizado.",
            "Conservar control de las remisiones emitidas y sus anulaciones.",
        ],
    )

    add_heading(document, "14. Cierre", level=1)
    add_paragraph(
        document,
        "Sistema IRIS ha sido concebido para apoyar la operacion interna de EPSI HL mediante una interfaz organizada y procesos estandarizados. El uso adecuado del sistema contribuye a mejorar la trazabilidad de la informacion, la seguridad operativa y la eficiencia del trabajo diario."
    )

    for section in document.sections:
        add_footer(section, "Manual de Usuario - Sistema IRIS - EPSI HL")

    return document


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    technical_path = OUTPUT_DIR / "Manual_Tecnico_Sistema_IRIS.docx"
    user_path = OUTPUT_DIR / "Manual_de_Usuario_Sistema_IRIS.docx"

    create_technical_manual().save(technical_path)
    create_user_manual().save(user_path)

    print(technical_path)
    print(user_path)


if __name__ == "__main__":
    main()
